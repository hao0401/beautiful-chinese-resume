#!/usr/bin/env python
"""Build a polished one-page Chinese resume DOCX from structured JSON."""

from __future__ import annotations

import argparse
import json
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "Missing dependency: python-docx. Install with `python -m pip install python-docx`."
    ) from exc


@dataclass(frozen=True)
class ResumeStyle:
    accent: str
    border: str
    body: str = "202124"
    muted: str = "6B7280"
    secondary: str = "4B5563"
    font: str = "Microsoft YaHei"
    name_size: float = 21
    role_size: float = 10.5
    contact_size: float = 9.2
    section_size: float = 11
    heading_size: float = 10
    body_size: float = 9.35
    margin_top_cm: float = 1.35
    margin_bottom_cm: float = 1.25
    margin_left_cm: float = 1.35
    margin_right_cm: float = 1.35


STYLE_PRESETS = {
    "campus": ResumeStyle(accent="2F6F73", border="D8E4E6"),
    "business": ResumeStyle(accent="31516A", border="DDE5EA", margin_left_cm=1.45, margin_right_cm=1.45),
    "fresh": ResumeStyle(accent="3C7A89", border="D7E8EC", role_size=10.3),
    "minimal": ResumeStyle(accent="404040", border="E2E2E2", section_size=10.8),
}

SECTION_ORDER = ["教育背景", "项目经历", "实习经历", "技能证书", "校园经历", "获奖经历", "自我评价"]


def load_json(path: str) -> dict[str, Any]:
    if path == "-":
        return json.loads(sys.stdin.read().lstrip("\ufeff"))
    with open(path, "r", encoding="utf-8-sig") as fh:
        return json.load(fh)


def set_run_font(
    run,
    *,
    font: str = "Microsoft YaHei",
    size: float | None = None,
    bold: bool = False,
    color: str | None = None,
) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 0, line: float = 1.0) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_bottom_border(paragraph, color: str = "D8E4E6") -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_text(
    paragraph,
    text: str,
    *,
    style: ResumeStyle,
    size: float = 10,
    bold: bool = False,
    color: str | None = None,
):
    run = paragraph.add_run(text)
    set_run_font(run, font=style.font, size=size, bold=bold, color=color)
    return run


def normalize_sections(data: dict[str, Any]) -> list[dict[str, Any]]:
    sections = data.get("sections")
    if isinstance(sections, list):
        return [s for s in sections if s.get("title")]

    mapped: list[dict[str, Any]] = []
    key_map = [
        ("education", "教育背景"),
        ("projects", "项目经历"),
        ("internships", "实习经历"),
        ("experience", "实习经历"),
        ("skills", "技能证书"),
        ("certificates", "技能证书"),
        ("campus", "校园经历"),
        ("awards", "获奖经历"),
    ]
    for key, title in key_map:
        value = data.get(key)
        if not value:
            continue
        if key in {"skills", "certificates"} and all(isinstance(v, str) for v in value):
            items = [{"heading": title, "details": ["；".join(value)]}]
        elif isinstance(value, list):
            items = value
        else:
            items = [{"heading": title, "details": [str(value)]}]

        existing = next((s for s in mapped if s["title"] == title), None)
        if existing:
            existing["items"].extend(items)
        else:
            mapped.append({"title": title, "items": items})

    return mapped


def sort_sections(sections: list[dict[str, Any]]) -> list[dict[str, Any]]:
    order = {title: idx for idx, title in enumerate(SECTION_ORDER)}

    def section_sort_key(section: dict[str, Any]) -> int:
        return order.get(section["title"], 99)

    return sorted(sections, key=section_sort_key)


def add_header(doc: Document, data: dict[str, Any], style: ResumeStyle) -> None:
    name = data.get("name") or "姓名"
    target_role = data.get("target_role") or data.get("role") or "求职方向待补充"
    target_company = data.get("target_company")
    show_company = bool(data.get("show_target_company"))
    contact = data.get("contact") or []
    if isinstance(contact, str):
        contact = [contact]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=1)
    add_text(p, name, style=style, size=style.name_size, bold=True, color=style.body)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=1)
    role_line = f"求职意向：{target_role}"
    if target_company and show_company:
        role_line += f" | 目标公司：{target_company}"
    add_text(p, role_line, style=style, size=style.role_size, bold=True, color=style.accent)

    if contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=5)
        add_text(
            p,
            " | ".join(str(item) for item in contact if item),
            style=style,
            size=style.contact_size,
            color="5F6368",
        )


def add_summary(doc: Document, summary: Any, style: ResumeStyle) -> None:
    if not summary:
        return
    if isinstance(summary, str):
        summary = [summary]
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=3)
    add_text(
        p,
        "；".join(str(item).strip() for item in summary if str(item).strip()),
        style=style,
        size=9.8,
        color="3C4043",
    )


def add_section_title(doc: Document, title: str, style: ResumeStyle) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=2)
    add_text(p, title, style=style, size=style.section_size, bold=True, color=style.accent)
    add_bottom_border(p, style.border)


def add_item(doc: Document, item: Any, style: ResumeStyle) -> None:
    if isinstance(item, str):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=1)
        add_text(p, "• " + item, style=style, size=9.7, color=style.body)
        return

    heading = item.get("heading") or item.get("title") or ""
    date = item.get("date") or item.get("time") or ""
    subheading = item.get("subheading") or item.get("subtitle") or ""
    details = item.get("details") or item.get("bullets") or []
    if isinstance(details, str):
        details = [details]

    if heading or date:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=1, after=0.5)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(18.0), WD_TAB_ALIGNMENT.RIGHT)
        add_text(p, str(heading), style=style, size=style.heading_size, bold=True, color=style.body)
        if date:
            add_text(p, "\t" + str(date), style=style, size=9.2, color=style.muted)

    if subheading:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=0.5)
        add_text(p, str(subheading), style=style, size=9.5, color=style.secondary)

    for detail in details:
        text = str(detail).strip()
        if not text:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.15)
        p.paragraph_format.first_line_indent = Cm(-0.15)
        set_paragraph_spacing(p, after=0.2)
        add_text(p, "• " + text, style=style, size=style.body_size, color=style.body)


def build_docx(data: dict[str, Any], output: Path, style_name: str) -> None:
    style = STYLE_PRESETS.get(style_name)
    if style is None:
        valid = ", ".join(sorted(STYLE_PRESETS))
        raise ValueError(f"unknown style {style_name!r}; valid styles: {valid}")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(style.margin_top_cm)
    section.bottom_margin = Cm(style.margin_bottom_cm)
    section.left_margin = Cm(style.margin_left_cm)
    section.right_margin = Cm(style.margin_right_cm)

    normal = doc.styles["Normal"]
    normal.font.name = style.font
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), style.font)
    normal.font.size = Pt(9.5)

    add_header(doc, data, style)
    add_summary(doc, data.get("summary"), style)

    for section_data in sort_sections(normalize_sections(data)):
        items = section_data.get("items") or []
        if not items:
            continue
        add_section_title(doc, section_data["title"], style)
        for item in items:
            add_item(doc, item, style)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_json", nargs="?", help="Resume JSON path, or '-' for stdin.")
    parser.add_argument("output_docx", nargs="?", help="Output .docx path.")
    parser.add_argument("--style", choices=sorted(STYLE_PRESETS), default="campus")
    parser.add_argument("--list-styles", action="store_true", help="List available style presets and exit.")
    args = parser.parse_args()

    if args.list_styles:
        print("\n".join(sorted(STYLE_PRESETS)))
        return 0

    if not args.input_json or not args.output_docx:
        parser.error("input_json and output_docx are required unless --list-styles is used")

    data = load_json(args.input_json)
    output = Path(args.output_docx)
    build_docx(data, output, args.style)
    print(f"wrote: {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
